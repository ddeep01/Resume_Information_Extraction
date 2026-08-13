from pathlib import Path

import fitz
import pdfplumber
from docx import Document


class PDFExtractor:

    MIN_TEXT_LENGTH = 100

    @staticmethod
    def extract_pymupdf(pdf_path):

        try:
            text_parts = []

            with fitz.open(pdf_path) as doc:

                for page in doc:
                    text_parts.append(page.get_text())

            return "\n".join(text_parts).strip()

        except Exception as e:
            print(f"PyMuPDF Error: {pdf_path} -> {e}")
            return ""

    @staticmethod
    def extract_pdfplumber(pdf_path):

        try:
            text_parts = []

            with pdfplumber.open(pdf_path) as pdf:

                for page in pdf.pages:

                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(page_text)

            return "\n".join(text_parts).strip()

        except Exception as e:
            print(f"pdfplumber Error: {pdf_path} -> {e}")
            return ""

    @classmethod
    def extract(cls, pdf_path):

        text = cls.extract_pymupdf(pdf_path)

        if len(text) >= cls.MIN_TEXT_LENGTH:
            return text, "pymupdf"

        text = cls.extract_pdfplumber(pdf_path)

        if len(text) >= cls.MIN_TEXT_LENGTH:
            return text, "pdfplumber"

        return "", "failed"


class DOCXExtractor:

    MIN_TEXT_LENGTH = 100

    @staticmethod
    def extract_docx(docx_path):

        try:
            document = Document(docx_path)

            text_parts = []



            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:
                    text_parts.append(text)


            for table in document.tables:

                for row in table.rows:

                    row_text = []

                    for cell in row.cells:

                        cell_text = cell.text.strip()

                        if cell_text:
                            row_text.append(cell_text)

                    if row_text:
                        text_parts.append(" | ".join(row_text))

            text = "\n".join(text_parts).strip()

            if len(text) >= DOCXExtractor.MIN_TEXT_LENGTH:
                return text, "python-docx"

            return "", "failed"

        except Exception as e:

            print(f"DOCX Error: {docx_path} -> {e}")

            return "", "failed"


RAW_DIR = Path("data/raw_resumes")
OUTPUT_DIR = Path("data/extracted_text")

OUTPUT_DIR.mkdir(
    parents=True,
    exist_ok=True
)


def process_all_resumes():


    files = [
        file
        for file in RAW_DIR.iterdir()
        if file.is_file()
        and file.suffix.lower() in [".pdf", ".docx"]
    ]

    print(f"\nFound {len(files)} PDF/DOCX files\n")

    success = 0
    failed = 0

    for file in files:

        extension = file.suffix.lower()


        if extension == ".pdf":

            text, method = PDFExtractor.extract(
                str(file)
            )

        elif extension == ".docx":

            text, method = DOCXExtractor.extract_docx(
                str(file)
            )


        else:

            print(
                f"[UNSUPPORTED] {file.name}"
            )

            continue


        if method == "failed":

            failed += 1

            print(
                f"[FAILED] {file.name}"
            )

            continue


        output_file = (
            OUTPUT_DIR /
            f"{file.stem}.txt"
        )

        output_file.write_text(
            text,
            encoding="utf-8"
        )

        success += 1

        print(
            f"[SUCCESS] {file.name} -> {method}"
        )

    print("\n" + "=" * 50)
    print(f"Total Files : {len(files)}")
    print(f"Success     : {success}")
    print(f"Failed      : {failed}")
    print("=" * 50)


if __name__ == "__main__":
    process_all_resumes()